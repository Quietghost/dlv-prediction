from docx import Document
import io
import shutil
import os
import re


def convertDocxToText(path):
    for d in os.listdir(path):
        fullText = ""
        fileExtension = d.split(".")[-1]
        if fileExtension == "docx":
            docxFilename = path + d
            document = Document(docxFilename)
            textFilename = path + d.split(".")[0] + ".txt"
            with io.open(textFilename, "w", encoding="utf-8") as textFile:
                for para in document.paragraphs:
                    fullText += " " + para.text
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            fullText += " " + cell.text
                noSpaces = re.sub(' +', ' ', fullText)
                noNewline = re.sub('\n', ' ', noSpaces)
                noTabs = re.sub('\t+', ' ', noNewline)
                textFile.write(noTabs)
            print(docxFilename)

path = "/home/dev/team_blue/python/dlvs_random/"
convertDocxToText(path)
