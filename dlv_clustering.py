from docx import Document
import io
import shutil
import os
import re
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sklearn.metrics import adjusted_rand_score
from sklearn import externals
from stop_words import get_stop_words
from sklearn import metrics
import csv
import sys
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.cm as cm
import numpy as np


def convertDocxToText(path):
    for d in os.listdir(path):
        fullText = ""
        fileExtension = d.split(".")[-1]
        if fileExtension == "docx":
            docxFilename = path + d
            print("Start: " + docxFilename)
            try:
                document = Document(docxFilename)
                textFilename = path + d.split(".")[0] + ".txt"
            except:
                errors.append(docxFilename)
                pass
            with io.open(textFilename, "w", encoding="utf-8") as textFile:
                try:
                    for para in document.paragraphs:
                        fullText += " " + para.text
                    for table in document.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                fullText += " " + cell.text
                    noSpaces = re.sub(' +', ' ', fullText)
                    noNewline = re.sub('\n+', ' ', noSpaces)
                    noTabs = re.sub('\t+', ' ', noNewline)
                    noComma = noTabs.replace(',', '')
                    noSemiColon = noComma.replace(';', '')
                    nothing = re.sub('[^A-Za-z0-9\xe4\xF6\xFC\n+\t+ +]+', '',
                                     noSemiColon)
                    textFile.write(nothing)
                    documents.append(nothing.encode('utf-8'))
                    filenames.append(docxFilename)
                    print("End: " + docxFilename)
                except:
                    errors.append(docxFilename)
                    pass
    for files in errors:
        print(files)


def getTxtFilenames(path):
    for d in os.listdir(path):
        fullText = ""
        fileExtension = d.split(".")[-1]
        if fileExtension == "txt":
            txtFilename = path + d
            print("Start: " + txtFilename)
            try:
                textFilename = path + d.split(".")[0] + ".txt"
                statinfo = os.stat(textFilename)
                if statinfo.st_size > 0:
                    filenames.append(txtFilename)
            except:
                errors.append(txtFilename)
                pass


def createCSV(documents):
    print(str(len(documents)))
    csvfile = "dataset_filenames.csv"
    try:
        with io.open(csvfile, "w", encoding="utf-8") as csvFile:
            n = 1
            csvFile.write(unicode("label_column,text_column\n"))
            for item in documents:
                csvFile.write(str(n) + "," + item.decode('utf-8') + "\n")
                n = n + 1
    except IOError:
        print("I/O error")


def importCSV(filename):
    df = pd.read_csv(filename)
    columns = df.text_column
    for item in columns:
        documents.append(item)
    print(documents[0])


def cluster(documents):
    sw = get_stop_words('german')

    vectorizer = TfidfVectorizer(stop_words=sw)
    X = vectorizer.fit_transform(documents)

    vectorizer_filename = 'dlv-vectorizer.joblib.z'
    externals.joblib.dump(vectorizer, vectorizer_filename)

    true_k = 5
    model = KMeans(n_clusters=true_k, init='k-means++',
                   max_iter=300, n_init=10)
    model.fit(X)

    labels = model.predict(X)

    model_filename = 'dlv-clusters.joblib.z'
    externals.joblib.dump(model, model_filename)

    labels_filename = 'dlv-labels.joblib.z'
    externals.joblib.dump(labels, labels_filename)

    print("\n")
    print(labels)

    clusters = {}
    n = 0
    for item in labels:
        if item in clusters:
            clusters[item].append(filenames[n])
        else:
            clusters[item] = [filenames[n]]
        n += 1

    for item in clusters:
        print("\n")
        print "Cluster ", item
        for i in clusters[item]:
            print i

    clusters_filename = 'dlv-clusteritems.joblib.z'
    externals.joblib.dump(clusters, clusters_filename)

    print("\n")
    print("Top terms per cluster:")
    order_centroids = model.cluster_centers_.argsort()[:, ::-1]
    terms = vectorizer.get_feature_names()
    for i in range(true_k):
        print("Cluster %d:" % i),
        for ind in order_centroids[i, :10]:
            print(' %s' % terms[ind]),
        print

    print("Homogeneity: %0.3f" % metrics.homogeneity_score(labels,
          model.labels_))
    print("Completeness: %0.3f" % metrics.completeness_score(labels,
          model.labels_))
    print("V-measure: %0.3f" % metrics.v_measure_score(labels, model.labels_))
    print("Adjusted Rand-Index: %.3f"
          % metrics.adjusted_rand_score(labels, model.labels_))
    print("Silhouette Coefficient: %0.3f"
          % metrics.silhouette_score(X, model.labels_, sample_size=1000))


def evaluation(documents):
    range_n_clusters = [2, 3, 4, 5, 6, 7, 8, 9, 10, 20, 25]

    for n_clusters in range_n_clusters:

        # Create a subplot with 1 row and 2 columns
        fig, (ax1, ax2) = plt.subplots(1, 2)
        fig.set_size_inches(18, 7)

        # The 1st subplot is the silhouette plot
        # The silhouette coefficient can range from -1, 1 but in t
        # his example all
        # lie within [-0.1, 1]
        ax1.set_xlim([-1, 1])
        # The (n_clusters+1)*10 is for inserting blank space between silhouette
        # plots of individual clusters, to demarcate them clearly.

        sw = get_stop_words('german')

        vectorizer = TfidfVectorizer(stop_words=sw)
        X = vectorizer.fit_transform(documents)

        ax1.set_ylim([0, len(X) + (n_clusters + 1) * 10])

        model = KMeans(n_clusters=n_clusters, init='k-means++',
                       max_iter=300, n_init=10)
        model.fit(X)

        labels = model.predict(X)

        silhouette_avg = metrics.silhouette_score(X, model.labels)

        print("For n_clusters =", n_clusters,
              "The average silhouette_score is :", silhouette_avg)

        y_lower = 10
        for i in range(n_clusters):
            # Aggregate the silhouette scores for samples belonging to
            # cluster i, and sort them
            ith_cluster_silhouette_values = \
                sample_silhouette_values[cluster_labels == i]

            ith_cluster_silhouette_values.sort()

            size_cluster_i = ith_cluster_silhouette_values.shape[0]
            y_upper = y_lower + size_cluster_i

            color = cm.nipy_spectral(float(i) / n_clusters)
            ax1.fill_betweenx(np.arange(y_lower, y_upper),
                              0, ith_cluster_silhouette_values,
                              facecolor=color, edgecolor=color, alpha=0.7)

            # Label the silhouette plots with their cluster numbers at
            # the middle
            ax1.text(-0.05, y_lower + 0.5 * size_cluster_i, str(i))

            # Compute the new y_lower for next plot
            y_lower = y_upper + 10  # 10 for the 0 samples

        ax1.set_title("The silhouette plot for the various clusters.")
        ax1.set_xlabel("The silhouette coefficient values")
        ax1.set_ylabel("Cluster label")

        # The vertical line for average silhouette score of all the values
        ax1.axvline(x=silhouette_avg, color="red", linestyle="--")

        ax1.set_yticks([])  # Clear the yaxis labels / ticks
        ax1.set_xticks([-0.1, 0, 0.2, 0.4, 0.6, 0.8, 1])


documents = []
filenames = []
errors = []
path = "/home/dev/team_blue/python/test/"
# convertDocxToText(path)
# createCSV(documents)
# getTxtFilenames(path)
importCSV("trainingset.csv")
# cluster(documents)
evaluation(documents)
plt.show()
